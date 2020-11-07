from docx import Document
import os.path
import documents_x
import unittest

class Tests_documents(unittest.TestCase):

    def test_1(self):
        document = Document()                               # создаем тестовый файл и наполним его текстом
        document.add_heading('Главный заголовок', 0)
        p = document.add_paragraph('В этом параграфе что-то есть')
        document.add_paragraph('и в этом параграфе тоже стоит какая-то фигня')
        document.add_paragraph('И в конце какой-то текст')
        document.save('testfile.docx')                      # перед вводом в функцию сохранимся
        
        name = 'testfile.docx'
        old = 'стоит'
        new = 'лежит'
        documents_x.my_replacement(name, old, new)
        document = Document(name)
        self.assertEqual(document.paragraphs[0].text, 'Главный заголовок')
        self.assertEqual(document.paragraphs[1].text, 'В этом параграфе что-то есть')
        self.assertEqual(document.paragraphs[2].text, 'и в этом параграфе тоже лежит какая-то фигня')
        self.assertEqual(document.paragraphs[3].text, 'И в конце какой-то текст')


        
    def tearDown(self):    
        os.remove('testfile.docx')                          # в конце удаляем тестовый файл
        
if __name__ == '__main__':
    try: unittest.main()
    except SystemExit: pass

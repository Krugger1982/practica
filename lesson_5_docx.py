from docx import Document
def my_replacement(name, bylo, stalo):
    document = Document(name)
    for paragraf in document.paragraphs:
        content = paragraf.text
        position = content.find(bylo)
        if position != -1:
            content = content[:position] + stalo + content[position + len(bylo):]
            paragraf.clear()
            paragraf.add_run(content)
    document.save(name)
